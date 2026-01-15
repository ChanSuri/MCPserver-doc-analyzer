import os
import sys
import logging
from typing import List, Dict, Optional
from docx import Document
from fastmcp import FastMCP
from rapidfuzz import fuzz

# --- Initialization ---
logging.basicConfig(level=logging.INFO, stream=sys.stderr)
logger = logging.getLogger(__name__)

mcp = FastMCP("Arduino_Analytics_Expert")
DOC_PATH = os.path.join(os.path.dirname(__file__), "Arduino - The Analytics Ecosytem playbook.docx")

_knowledge_index = None
_last_mtime = None


# --- Analysis functions ---
def build_knowledge_index():
    """Make index with Heading"""
    global _knowledge_index, _last_mtime

    if not os.path.exists(DOC_PATH):
        logger.error(f"Document not found at {DOC_PATH}")
        return []
    
    # check file modification time
    mtime = os.path.getmtime(DOC_PATH)
    if _knowledge_index is not None and _last_mtime == mtime:
        return _knowledge_index  # cache valid

    _last_mtime = mtime
    _knowledge_index = []

    doc = Document(DOC_PATH)
    index = []
    current_main_topic = "General Overview"
    current_sub_topic = ""
    current_content = []

    # snapshot current section
    def save_current_section():
        if current_content:
            text = "\n".join(current_content).strip()
            if text:
                index.append({
                    "main_topic": current_main_topic,
                    "sub_topic": current_sub_topic or current_main_topic,
                    "content": text,
                    "search_key": f"{current_main_topic} {current_sub_topic}".lower(),
                })

    for child in doc.element.body.iterchildren():
        # <p> paragraph
        if child.tag.endswith('p'):
            para = [p for p in doc.paragraphs if p._element == child][0]
            text = para.text.strip()
            if not text: continue
            
            is_bold = para.runs[0].bold if para.runs else False
            # Big title (Heading 1)
            if para.style.name.startswith('Heading 1'):
                save_current_section()
                current_main_topic = text
                current_sub_topic = ""
                current_content = []
            # Small title (Heading 2 or bold short line)
            elif para.style.name.startswith('Heading') or (is_bold and len(text) < 60):
                save_current_section()
                current_sub_topic = text
                current_content = []
            else:
                current_content.append(text)
        
        # <tbl> table
        elif child.tag.endswith('tbl'):
            table = [t for t in doc.tables if t._element == child][0]
            current_content.append(parse_table_to_markdown(table))

        # graph
        elif child.findall('.//w:drawing', namespaces=doc.element.nsmap):
            img_context = current_sub_topic or "Unlabeled Section"
            current_content.append(f"\n> [Graph in this section: {img_context}] (Please search in original document)\n")

    save_current_section()
    _knowledge_index = index
    return _knowledge_index


def smart_search(query: str, main_filter: Optional[str] = None, top_k: int = 3) -> List[Dict]:
    """
    Search logic returning internal Dictionary structure.
    Increased top_k slightly to ensure context isn't missed.
    """
    kb = build_knowledge_index()
    results = []
    
    for item in kb:
        if main_filter and main_filter.lower() not in item['main_topic'].lower():
            continue
            
        # score: title(70%) + content(30%)
        t_score = fuzz.WRatio(query, item['sub_topic'])
        c_score = fuzz.partial_ratio(query.lower(), item['content'].lower())
        final_score = t_score * 0.7 + c_score * 0.3
        
        if final_score > 40:
            results.append({**item, "score": final_score})
            
    return sorted(results, key=lambda x: x['score'], reverse=True)[:top_k] 


# --- Helper: Format Output as Markdown ---
def format_results_as_markdown(title: str, results: List[Dict], error_msg: str = "No info found.") -> str:
    """
    Converts search results into a clean Markdown string.
    This reduces token usage and processing time for the LLM compared to JSON.
    """
    if not results:
        return f"**Status**: {error_msg}\nQuery: {title}"

    md_lines = [f"# Search Results for: '{title}'\n"]
    for i, r in enumerate(results, 1):
        md_lines.append(f"## {i}. {r['main_topic']} > {r['sub_topic']}")
        
        # Truncate very long content if necessary to prevent timeouts
        # 2000 chars is roughly 500-800 tokens, which is a safe chunk size per section
        content = r['content']
        if len(content) > 2000: 
            content = content[:2000] + "\n...(content truncated due to length)..."
        
        md_lines.append(f"{content}\n")
        md_lines.append("---")
    
    return "\n".join(md_lines)

def parse_table_to_markdown(table) -> str:
    """change Word table to Markdown format for LLM"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0: # header separator
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return "\n" + "\n".join(rows) + "\n"


# --- MCP Tools ---

@mcp.tool()
async def get_comprehensive_overview() -> str:
    """
    Provides a structural overview of the entire Analytics Ecosystem.
    Returns a Markdown formatted list.
    """
    kb = build_knowledge_index()
    structure = {}
    for item in kb:
        structure.setdefault(item['main_topic'], []).append(item['sub_topic'])
    
    md_lines = ["# Analytics Ecosystem Overview\n"]
    
    for main, subs in structure.items():
        unique_subs = sorted(list(set(subs)))
        md_lines.append(f"### {main}")
        for sub in unique_subs:
            # Avoid repeating if subtopic is same as main topic
            if sub and sub != main:
                md_lines.append(f"- {sub}")
        md_lines.append("") # Empty line for spacing
        
    return "\n".join(md_lines)


@mcp.tool()
async def solve_analytics_issue(query: str) -> str:
    """
    Search for solutions to specific problems, including discrepancies or implementation errors.
    """
    results = smart_search(query)
    return format_results_as_markdown(query, results, "No matching solution found in playbook.")


@mcp.tool()
async def check_limits_and_compliance(topic: str) -> str:
    """
    Query GA4 limits, data retention, cookie consent, or age restrictions.
    """
    results = smart_search(topic, main_filter="restrictions")
    results += smart_search(topic, main_filter="limits")
    
    # Remove duplicates based on content
    unique_results = {r['content']: r for r in results}.values()
    
    return format_results_as_markdown(topic, list(unique_results), "No compliance or limit information found.")


@mcp.tool()
async def compare_platform_strategy(feature_or_tool: str) -> str:
    """
    Technical comparison between platforms (GA4, Segment, Shopify) and platform choice strategy.
    """
    results = smart_search(feature_or_tool, main_filter="choose")
    results += smart_search(feature_or_tool, main_filter="discrepancies")
    
    # Remove duplicates
    unique_results = {r['content']: r for r in results}.values()

    return format_results_as_markdown(feature_or_tool, list(unique_results), "No strategic comparison found.")


@mcp.tool()
async def get_metric_definition(term: str) -> str:
    """
    Look up the precise definition of a specific metric or term (e.g., 'Session', 'Attribution Window').
    """
    # Search in "Dimensions and Metrics" section first
    results = smart_search(term, main_filter="Dimensions and Metrics")

    # If not found, try searching the full document for "Definition"
    if not results:
        results = smart_search(term + " definition")

    # Only return the top 1 result for definitions to keep it concise
    return format_results_as_markdown(term, results[:1], "Definition not found in the glossary.")


@mcp.tool()
async def report_documentation_issue(section_topic: str, issue_description: str) -> str:
    """
    Allow employees to report errors or outdated information in the playbook.
    """
    log_entry = f"[REPORT] Topic: {section_topic} | Issue: {issue_description}"
    logger.warning(log_entry)
    
    # Append to a local log file for review by Data Governance team
    try:
        with open("playbook_feedback.log", "a", encoding="utf-8") as f:
            f.write(log_entry + "\n")
    except Exception as e:
        logger.error(f"Could not write to log file: {e}")
        
    return f"Feedback logged successfully for topic: '{section_topic}'."


# --- start ---

def main():
    logger.info("Arduino Analytics MCP Server is initializing...")
    try:
        kb = build_knowledge_index()
        logger.info(f"Successfully indexed {len(kb)} sections from Playbook.")
    except Exception as e:
        logger.error(f"Failed to load Playbook: {e}")
        sys.exit(1)

    mcp.run(transport="stdio")

if __name__ == "__main__":
    main()